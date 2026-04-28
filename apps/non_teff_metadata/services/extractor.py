"""
Non-TEFF Metadata Extractor — multi-format extraction service.

Each extractor returns a list of dicts matching the field schema defined in
config/non_teff_fields.json.  Regex patterns are defined as module-level
constants so they can be tuned without diving into logic code.
"""

import logging
import re
import threading

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# SOFT-CODED field regex patterns (module-level — not inline magic strings)
# ---------------------------------------------------------------------------

# Patterns for instrument tag numbers (e.g. FIC-101, PT-2001A, LT_305)
INSTRUMENT_TAG_PATTERN = re.compile(
    r'\b([A-Z]{1,4}[-_][0-9]{2,5}[A-Z]?)\b'
)

# Line numbers (e.g. 6"-P-1001-A1A, 4"-GAS-2005)
LINE_NUMBER_PATTERN = re.compile(
    r'\b(\d{1,4}"[-–]\w{1,6}[-–]\d{3,6}(?:[-–]\w{1,6})?)\b'
)

# Equipment numbers (e.g. E-101, V-201, P-301A, HE-4002)
EQUIPMENT_NO_PATTERN = re.compile(
    r'\b([A-Z]{1,3}-\d{3,5}[A-Z]?)\b'
)

# Document/drawing numbers (e.g. P16093-PR-PFD-001, RAD-MECH-DS-003)
DOCUMENT_NO_PATTERN = re.compile(
    r'\b([A-Z0-9]{2,8}-[A-Z]{2,8}-[A-Z]{2,4}-\d{3,5}(?:-\d{2,3})?)\b'
)

# Revision strings (e.g. "Rev. A", "Rev 03", "Revision B")
REVISION_PATTERN = re.compile(
    r'\b[Rr]ev(?:ision)?\.?\s*([A-Z0-9]{1,3})\b'
)

# Date patterns (ISO and common engineering formats)
DATE_PATTERN = re.compile(
    r'\b(\d{4}-\d{2}-\d{2}|\d{2}[./]\d{2}[./]\d{4}|\d{2}-[A-Z]{3}-\d{4})\b'
)

# Discipline keywords
DISCIPLINE_KEYWORDS = {
    'process': 'Process',
    'instrument': 'Instrument',
    'electrical': 'Electrical',
    'mechanical': 'Mechanical',
    'piping': 'Piping',
    'civil': 'Civil',
    'structural': 'Structural',
    'hvac': 'HVAC',
    'safety': 'Safety',
}

# Mechanical component keywords
MECHANICAL_COMPONENT_KEYWORDS = [
    'pump', 'compressor', 'vessel', 'heat exchanger', 'separator',
    'filter', 'valve', 'column', 'tank', 'drum', 'reactor', 'agitator',
    'blower', 'fan', 'turbine', 'generator', 'motor', 'gearbox',
]

# Status keywords for Non-TEFF context
NON_TEFF_STATUS_KEYWORDS = [
    'non-teff', 'non teff', 'preliminary', 'draft', 'issued for comment',
    'issued for approval', 'ifr', 'ifa', 'ifc', 'afc', 'issued for construction',
    'for information', 'for review',
]


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _first_match(pattern, text):
    """Return first captured group from pattern search, or empty string."""
    m = pattern.search(text)
    return m.group(1) if m else ''


def _all_matches(pattern, text):
    """Return all unique matches from pattern findall, joined by ', '."""
    matches = list(dict.fromkeys(pattern.findall(text)))  # deduplicate, preserve order
    return ', '.join(matches[:5])  # cap at 5 to keep cells readable


def _detect_discipline(text):
    text_lower = text.lower()
    for keyword, label in DISCIPLINE_KEYWORDS.items():
        if keyword in text_lower:
            return label
    return ''


def _detect_mechanical_component(text):
    text_lower = text.lower()
    found = [k.title() for k in MECHANICAL_COMPONENT_KEYWORDS if k in text_lower]
    return ', '.join(list(dict.fromkeys(found))[:3])


def _detect_status(text):
    text_lower = text.lower()
    for kw in NON_TEFF_STATUS_KEYWORDS:
        if kw in text_lower:
            return kw.upper()
    return ''


def _extract_fields_from_text(text, source_label=''):
    """
    Run all field extractors against a block of text.
    Returns a single metadata dict.
    """
    return {
        'document_no': _all_matches(DOCUMENT_NO_PATTERN, text),
        'document_title': '',         # filled from structured sources below
        'revision': _first_match(REVISION_PATTERN, text),
        'discipline': _detect_discipline(text),
        'instrument_tag_no': _all_matches(INSTRUMENT_TAG_PATTERN, text),
        'line_number': _all_matches(LINE_NUMBER_PATTERN, text),
        'equipment_no': _all_matches(EQUIPMENT_NO_PATTERN, text),
        'mechanical_component': _detect_mechanical_component(text),
        'status': _detect_status(text),
        'date': _first_match(DATE_PATTERN, text),
        'originator': '',             # hard to detect generically from raw text
        'remarks': source_label,
    }


def _merge_results(results):
    """
    Merge a list of per-page/per-section dicts into a single summary dict.
    Multi-value fields are deduplicated and joined.
    """
    merged = {k: set() for k in results[0]} if results else {}
    scalar_keys = {'revision', 'date', 'originator', 'document_title', 'remarks'}

    for r in results:
        for k, v in r.items():
            if v:
                merged[k].add(v)

    return {
        k: list(v)[0] if k in scalar_keys else ', '.join(v)
        for k, v in merged.items()
    }


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------

def extract_from_pdf(file_path):
    """Extract metadata from a PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        return [{'error': 'pdfplumber not installed', 'remarks': 'PDF extraction unavailable'}]

    results = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ''
                if not text.strip():
                    continue
                rec = _extract_fields_from_text(text, source_label=f'Page {i + 1}')
                results.append(rec)
    except Exception as exc:
        logger.exception('PDF extraction error for %s', file_path)
        return [{'error': str(exc), 'remarks': 'PDF parsing failed'}]

    if not results:
        return [_extract_fields_from_text('', 'No text found')]

    return [_merge_results(results)]


def extract_from_excel(file_path):
    """Extract metadata rows from an Excel file using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        return [{'error': 'openpyxl not installed', 'remarks': 'Excel extraction unavailable'}]

    # Normalise column header → field key
    HEADER_MAP = {
        'document no': 'document_no',
        'document number': 'document_no',
        'doc no': 'document_no',
        'doc. no': 'document_no',
        'doc. no.': 'document_no',
        'document title': 'document_title',
        'title': 'document_title',
        'rev': 'revision',
        'revision': 'revision',
        'discipline': 'discipline',
        'instrument tag': 'instrument_tag_no',
        'instrument tag no': 'instrument_tag_no',
        'tag no': 'instrument_tag_no',
        'tag number': 'instrument_tag_no',
        'line number': 'line_number',
        'line no': 'line_number',
        'line no.': 'line_number',
        'equipment no': 'equipment_no',
        'equipment number': 'equipment_no',
        'equipment no.': 'equipment_no',
        'mechanical component': 'mechanical_component',
        'component': 'mechanical_component',
        'status': 'status',
        'date': 'date',
        'originator': 'originator',
        'prepared by': 'originator',
        'remarks': 'remarks',
        'comment': 'remarks',
    }

    results = []
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            headers = []
            for row in ws.iter_rows(min_row=1, max_row=50, values_only=True):
                # Detect header row: at least 3 non-None cells that match known headers
                row_lower = [str(c).strip().lower() if c is not None else '' for c in row]
                matched = sum(1 for h in row_lower if h in HEADER_MAP)
                if matched >= 2:
                    headers = [HEADER_MAP.get(h, None) for h in row_lower]
                    break

            if not headers:
                # No structured headers — fall back to full-sheet text scan
                all_text = ' '.join(
                    str(c) for r in ws.iter_rows(values_only=True) for c in r if c
                )
                rec = _extract_fields_from_text(all_text, ws.title)
                results.append(rec)
                continue

            # Map each data row to a result dict
            for row in ws.iter_rows(values_only=True):
                rec = {k: '' for k in [
                    'document_no', 'document_title', 'revision', 'discipline',
                    'instrument_tag_no', 'line_number', 'equipment_no',
                    'mechanical_component', 'status', 'date', 'originator', 'remarks',
                ]}
                has_data = False
                for idx, field_key in enumerate(headers):
                    if field_key and idx < len(row) and row[idx] is not None:
                        rec[field_key] = str(row[idx]).strip()
                        has_data = True
                if has_data:
                    results.append(rec)

    except Exception as exc:
        logger.exception('Excel extraction error for %s', file_path)
        return [{'error': str(exc), 'remarks': 'Excel parsing failed'}]

    return results if results else [_extract_fields_from_text('', 'No data found')]


def extract_from_word(file_path):
    """Extract metadata from a Word (.docx) file using python-docx."""
    try:
        import docx
    except ImportError:
        return [{'error': 'python-docx not installed', 'remarks': 'Word extraction unavailable'}]

    results = []
    try:
        doc = docx.Document(file_path)

        # Process tables first (more structured)
        for table in doc.tables:
            header_row = None
            for i, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                cells_lower = [c.lower() for c in cells]
                # Detect header row
                if any(kw in cells_lower for kw in ['document no', 'doc no', 'rev', 'revision', 'tag no']):
                    header_row = cells_lower
                    continue
                if header_row:
                    rec = _extract_fields_from_text(' '.join(cells), f'Table row {i}')
                    results.append(rec)

        # Process paragraph text as a bulk scan
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        if full_text.strip():
            rec = _extract_fields_from_text(full_text, 'Document body')
            results.append(rec)

    except Exception as exc:
        logger.exception('Word extraction error for %s', file_path)
        return [{'error': str(exc), 'remarks': 'Word parsing failed'}]

    if not results:
        return [_extract_fields_from_text('', 'No content found')]

    return [_merge_results(results)]


def extract_from_autocad(file_path):
    """Stub: AutoCAD DWG/DXF parsing not yet supported."""
    return [{
        'document_no': '',
        'document_title': '',
        'revision': '',
        'discipline': '',
        'instrument_tag_no': '',
        'line_number': '',
        'equipment_no': '',
        'mechanical_component': '',
        'status': '',
        'date': '',
        'originator': '',
        'remarks': 'AutoCAD (DWG/DXF) text extraction coming in a future release',
    }]


# ---------------------------------------------------------------------------
# Dispatcher — called from a background thread in views.py
# ---------------------------------------------------------------------------

FORMAT_DISPATCH = {
    'pdf': extract_from_pdf,
    'excel': extract_from_excel,
    'word': extract_from_word,
    'autocad': extract_from_autocad,
}


def dispatch_extraction(job_id, file_path, file_format):
    """
    Run extraction for the given job in the current thread.
    Updates the NonTeffExtractionJob model directly.

    This function is designed to be called from threading.Thread(target=...).
    """
    # Import here to avoid circular imports at module load time
    from apps.non_teff_metadata.models import NonTeffExtractionJob

    try:
        job = NonTeffExtractionJob.objects.get(pk=job_id)
        job.status = NonTeffExtractionJob.STATUS_PROCESSING
        job.progress = 10
        job.status_message = 'Extraction started…'
        job.save(update_fields=['status', 'progress', 'status_message'])

        extractor_fn = FORMAT_DISPATCH.get(file_format, extract_from_pdf)
        job.progress = 40
        job.status_message = 'Analysing document…'
        job.save(update_fields=['progress', 'status_message'])

        items = extractor_fn(file_path)

        job.progress = 90
        job.status_message = 'Finalising results…'
        job.save(update_fields=['progress', 'status_message'])

        job.result_json = {'items': items, 'total': len(items)}
        job.status = NonTeffExtractionJob.STATUS_COMPLETED
        job.progress = 100
        job.status_message = f'Extraction complete — {len(items)} record(s) found'
        job.save(update_fields=['result_json', 'status', 'progress', 'status_message'])

    except Exception as exc:
        logger.exception('dispatch_extraction failed for job %s', job_id)
        try:
            job = NonTeffExtractionJob.objects.get(pk=job_id)
            job.status = NonTeffExtractionJob.STATUS_FAILED
            job.error_message = str(exc)
            job.status_message = 'Extraction failed'
            job.progress = 0
            job.save(update_fields=['status', 'error_message', 'status_message', 'progress'])
        except Exception:
            pass


def run_extraction_async(job_id, file_path, file_format):
    """Spawn a daemon thread to run extraction without blocking the request."""
    t = threading.Thread(
        target=dispatch_extraction,
        args=(job_id, file_path, file_format),
        daemon=True,
    )
    t.start()
