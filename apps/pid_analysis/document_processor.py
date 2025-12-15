"""
Document Processing Service for RAG
Handles text extraction, chunking, and embedding generation
"""
import os
import io
from typing import List, Dict, Any
from django.conf import settings
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process documents for RAG system"""
    
    def __init__(self):
        """Initialize document processor"""
        self.chunk_size = int(os.getenv('RAG_CHUNK_SIZE', '1000'))
        self.chunk_overlap = int(os.getenv('RAG_CHUNK_OVERLAP', '200'))
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """
        Extract text from PDF file
        
        Args:
            pdf_file: File object or path
        
        Returns:
            Extracted text
        """
        try:
            import fitz  # PyMuPDF
            
            # Handle file object or path
            if isinstance(pdf_file, str):
                pdf_document = fitz.open(pdf_file)
            else:
                pdf_file.seek(0)
                pdf_bytes = pdf_file.read()
                pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            text = ""
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text += page.get_text()
            
            pdf_document.close()
            
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    def extract_text_from_docx(self, docx_file) -> str:
        """
        Extract text from DOCX file
        
        Args:
            docx_file: File object or path
        
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            # Handle file object or path
            if isinstance(docx_file, str):
                doc = Document(docx_file)
            else:
                docx_file.seek(0)
                doc = Document(docx_file)
            
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
            
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def extract_text_from_txt(self, txt_file) -> str:
        """
        Extract text from TXT file
        
        Args:
            txt_file: File object or path
        
        Returns:
            Text content
        """
        try:
            if isinstance(txt_file, str):
                with open(txt_file, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                txt_file.seek(0)
                text = txt_file.read().decode('utf-8')
            
            logger.info(f"Extracted {len(text)} characters from TXT")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise
    
    def extract_text(self, file_obj, filename: str) -> str:
        """
        Extract text from file based on extension
        
        Args:
            file_obj: File object
            filename: Original filename
        
        Returns:
            Extracted text
        """
        extension = filename.lower().split('.')[-1]
        
        if extension == 'pdf':
            return self.extract_text_from_pdf(file_obj)
        elif extension in ['docx', 'doc']:
            return self.extract_text_from_docx(file_obj)
        elif extension == 'txt':
            return self.extract_text_from_txt(file_obj)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks for embedding
        
        Args:
            text: Text to chunk
            metadata: Base metadata for all chunks
        
        Returns:
            List of chunks with metadata
        """
        chunks = []
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        chunk_num = 0
        
        for para in paragraphs:
            # If adding this paragraph exceeds chunk size
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append({
                    'text': current_chunk.strip(),
                    'metadata': {
                        **(metadata or {}),
                        'chunk_index': chunk_num,
                        'chunk_size': len(current_chunk)
                    }
                })
                chunk_num += 1
                
                # Start new chunk with overlap
                if self.chunk_overlap > 0:
                    # Take last N characters for overlap
                    overlap_text = current_chunk[-self.chunk_overlap:]
                    current_chunk = overlap_text + "\n" + para
                else:
                    current_chunk = para
            else:
                # Add to current chunk
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'metadata': {
                    **(metadata or {}),
                    'chunk_index': chunk_num,
                    'chunk_size': len(current_chunk)
                }
            })
        
        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def process_document(self, file_obj, filename: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Complete document processing pipeline
        
        Args:
            file_obj: File object
            filename: Original filename
            metadata: Document metadata
        
        Returns:
            List of processed chunks ready for embedding
        """
        # Extract text
        text = self.extract_text(file_obj, filename)
        
        # Clean text
        text = self.clean_text(text)
        
        # Chunk text
        chunks = self.chunk_text(text, metadata)
        
        return chunks
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text
        
        Args:
            text: Raw text
        
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]
        text = '\n'.join(lines)
        
        # Remove page numbers (simple heuristic)
        import re
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        
        # Normalize spaces
        text = re.sub(r' +', ' ', text)
        
        return text
