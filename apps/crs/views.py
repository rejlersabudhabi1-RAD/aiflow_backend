from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.db import transaction
from django.core.files.storage import default_storage
from django.http import HttpResponse
from django.conf import settings
from io import BytesIO
import json
import os

from .models import CRSDocument, CRSComment, CRSActivity, GoogleSheetConfig
from .serializers import (
    CRSDocumentSerializer,
    CRSDocumentDetailSerializer,
    CRSCommentSerializer,
    CRSActivitySerializer,
    GoogleSheetConfigSerializer,
    PDFExtractRequestSerializer,
    GoogleSheetExportSerializer,
    CRSCommentBulkCreateSerializer
)
from .pdf_extractor import PDFCommentExtractor
from .google_sheets_service import GoogleSheetsService

# Import helpers from crs_documents for PDF processing
try:
    from apps.crs_documents.helpers.comment_extractor import extract_reviewer_comments, get_comment_statistics
    from apps.crs_documents.helpers.template_populator import populate_crs_template
    from apps.crs_documents.helpers.template_manager import get_crs_template, get_template_info
    HELPERS_AVAILABLE = True
except ImportError:
    HELPERS_AVAILABLE = False


class CRSDocumentViewSet(viewsets.ModelViewSet):
    """ViewSet for CRS Document management"""
    queryset = CRSDocument.objects.all()
    permission_classes = [IsAuthenticated]
    
    def get_permissions(self):
        """
        Enhanced permission handling to ensure JWT authentication works properly
        """
        # For all actions, require authentication
        return [IsAuthenticated()]
    
    def get_serializer_class(self):
        if self.action == 'retrieve':
            return CRSDocumentDetailSerializer
        return CRSDocumentSerializer
    
    def perform_create(self, serializer):
        """Set uploaded_by to current user"""
        serializer.save(uploaded_by=self.request.user)
        
        # Log activity
        document = serializer.instance
        CRSActivity.objects.create(
            document=document,
            action='created',
            description=f'Document "{document.document_name}" created',
            performed_by=self.request.user
        )
    
    def perform_update(self, serializer):
        """Log update activity"""
        old_instance = self.get_object()
        old_status = old_instance.status
        
        serializer.save()
        
        new_instance = serializer.instance
        
        # Log status change
        if old_status != new_instance.status:
            CRSActivity.objects.create(
                document=new_instance,
                action='status_changed',
                description=f'Status changed from {old_status} to {new_instance.status}',
                performed_by=self.request.user,
                old_value={'status': old_status},
                new_value={'status': new_instance.status}
            )
        else:
            CRSActivity.objects.create(
                document=new_instance,
                action='updated',
                description=f'Document "{new_instance.document_name}" updated',
                performed_by=self.request.user
            )
    
    @action(detail=True, methods=['post'])
    def extract_pdf_comments(self, request, pk=None):
        """
        Extract comments from uploaded PDF file
        POST /api/v1/crs/documents/{id}/extract_pdf_comments/
        Body: { "auto_create_comments": true, "debug_mode": false }
        """
        document = self.get_object()
        
        serializer = PDFExtractRequestSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        if not document.pdf_file:
            return Response(
                {"error": "No PDF file uploaded for this document"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Get PDF file path
            pdf_path = document.pdf_file.path
            
            # Initialize extractor
            debug_mode = serializer.validated_data.get('debug_mode', False)
            extractor = PDFCommentExtractor(debug_mode=debug_mode)
            
            # Extract comments
            raw_comments = extractor.extract_comments_from_pdf(pdf_path)
            
            # Process comments
            processed_comments = extractor.process_extracted_comments(raw_comments)
            
            # Auto-create comments if requested
            auto_create = serializer.validated_data.get('auto_create_comments', True)
            
            if auto_create:
                with transaction.atomic():
                    # Delete existing comments for this document
                    document.comments.all().delete()
                    
                    # Create new comments
                    created_comments = []
                    for idx, comment_data in enumerate(processed_comments, start=1):
                        comment = CRSComment.objects.create(
                            document=document,
                            serial_number=idx,
                            page_number=comment_data['page'],
                            clause_number=comment_data.get('clause', ''),
                            comment_text=comment_data['text'],
                            comment_type=comment_data.get('type', 'other'),
                            color_rgb=comment_data.get('color'),
                            bbox=comment_data.get('bbox'),
                            status='open'
                        )
                        created_comments.append(comment)
                    
                    # Update document stats
                    document.total_comments = len(created_comments)
                    document.pending_comments = len(created_comments)
                    document.resolved_comments = 0
                    document.status = 'processing'
                    document.processed_at = timezone.now()
                    document.save()
                    
                    # Log activity
                    CRSActivity.objects.create(
                        document=document,
                        action='imported',
                        description=f'Imported {len(created_comments)} comments from PDF',
                        performed_by=request.user,
                        new_value={'comment_count': len(created_comments)}
                    )
            
            return Response({
                "success": True,
                "message": f"Extracted {len(processed_comments)} comments from PDF",
                "data": {
                    "total_comments": len(processed_comments),
                    "red_comments": sum(1 for c in processed_comments if c['type'] == 'red_comment'),
                    "yellow_boxes": sum(1 for c in processed_comments if c['type'] == 'yellow_box'),
                    "other_annotations": sum(1 for c in processed_comments if c['type'] not in ['red_comment', 'yellow_box']),
                    "comments": processed_comments[:10],  # Return first 10 for preview
                    "auto_created": auto_create
                }
            }, status=status.HTTP_200_OK)
        
        except FileNotFoundError as e:
            return Response(
                {"error": f"PDF file not found: {str(e)}"},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response(
                {"error": f"Error extracting comments: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['post'])
    def export_to_google_sheets(self, request, pk=None):
        """
        Export CRS data to Google Sheets
        POST /api/v1/crs/documents/{id}/export_to_google_sheets/
        Body: { "sheet_id": "optional_sheet_id", "start_row": 9, "auto_export": true }
        """
        document = self.get_object()
        
        serializer = GoogleSheetExportSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        # Get active Google Sheet config
        try:
            config = GoogleSheetConfig.objects.filter(is_active=True).first()
            if not config:
                return Response(
                    {"error": "No active Google Sheets configuration found. Please configure Google Sheets API first."},
                    status=status.HTTP_400_BAD_REQUEST
                )
        except Exception as e:
            return Response(
                {"error": f"Error loading Google Sheets configuration: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        # Get sheet ID
        sheet_id = serializer.validated_data.get('sheet_id')
        if not sheet_id:
            sheet_id = document.google_sheet_id
        
        if not sheet_id:
            return Response(
                {"error": "No Google Sheet ID provided and document has no linked sheet"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Initialize Google Sheets service
            gs_service = GoogleSheetsService(
                credentials_json=config.credentials_json,
                token_json=config.token_json
            )
            
            # Authenticate
            if not gs_service.authenticate():
                return Response(
                    {"error": "Failed to authenticate with Google Sheets API"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
            # Get comments
            comments = document.comments.all().order_by('serial_number')
            
            if not comments.exists():
                return Response(
                    {"error": "No comments to export"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Prepare data
            export_data = []
            for comment in comments:
                export_data.append({
                    'page': comment.page_number,
                    'clause': comment.clause_number or '',
                    'text': comment.comment_text,
                    'contractor_response': comment.contractor_response or '',
                    'company_response': comment.company_response or ''
                })
            
            # Export to sheet
            start_row = serializer.validated_data.get('start_row', 9)
            success, result = gs_service.export_to_sheet(
                spreadsheet_id=sheet_id,
                data=export_data,
                start_row=start_row
            )
            
            if success:
                # Update document
                document.google_sheet_id = sheet_id
                document.google_sheet_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}"
                document.save()
                
                # Save updated token if available
                new_token = gs_service.get_token_json()
                if new_token and new_token != config.token_json:
                    config.token_json = new_token
                    config.save()
                
                # Log activity
                CRSActivity.objects.create(
                    document=document,
                    action='exported',
                    description=f'Exported {len(export_data)} comments to Google Sheets',
                    performed_by=request.user,
                    new_value={'sheet_id': sheet_id, 'rows_exported': len(export_data)}
                )
                
                return Response({
                    "success": True,
                    "message": f"Successfully exported {len(export_data)} comments to Google Sheets",
                    "data": result
                }, status=status.HTTP_200_OK)
            else:
                return Response(
                    {"error": f"Export failed: {result.get('error', 'Unknown error')}"},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
        
        except Exception as e:
            return Response(
                {"error": f"Error exporting to Google Sheets: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=False, methods=['get'], permission_classes=[IsAuthenticated])
    def statistics(self, request):
        """
        Get CRS statistics
        GET /api/v1/crs/documents/statistics/
        """
        try:
            total_documents = CRSDocument.objects.count()
            total_comments = CRSComment.objects.count()
            
            documents_by_status = {}
            for choice in CRSDocument.STATUS_CHOICES:
                documents_by_status[choice[0]] = CRSDocument.objects.filter(status=choice[0]).count()
            
            comments_by_status = {}
            for choice in CRSComment.STATUS_CHOICES:
                comments_by_status[choice[0]] = CRSComment.objects.filter(status=choice[0]).count()
            
            return Response({
                "total_documents": total_documents,
                "total_comments": total_comments,
                "documents_by_status": documents_by_status,
                "comments_by_status": comments_by_status,
                "recent_documents": CRSDocumentSerializer(
                    CRSDocument.objects.all()[:5],
                    many=True,
                    context={'request': request}
                ).data
            })
        except Exception as e:
            return Response({
                'error': f'Statistics error: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=False, methods=['get'], url_path='config')
    def get_config(self, request):
        """
        Get CRS configuration (departments, etc.)
        GET /api/v1/crs/documents/config/
        
        Returns soft-coded configuration for frontend forms
        """
        try:
            # Load configuration from crs_config.json
            config_path = os.path.join(
                settings.BASE_DIR, 
                'apps', 'crs_documents', 'config', 'crs_config.json'
            )
            
            config_data = {
                'departments': [
                    {'id': 'process_control_hse', 'name': 'Process Control and HSE', 'code': 'PCH'},
                    {'id': 'ict', 'name': 'I&CT', 'code': 'ICT'},
                    {'id': 'structure_civil', 'name': 'Structure and Civil', 'code': 'SC'},
                ],
                'available_formats': ['xlsx', 'csv', 'json', 'pdf', 'docx'],
            }
            
            # Try to load from config file
            if os.path.exists(config_path):
                try:
                    with open(config_path, 'r') as f:
                        file_config = json.load(f)
                        if 'departments' in file_config:
                            config_data['departments'] = file_config['departments']
                except Exception:
                    pass  # Use defaults
            
            return Response({
                'success': True,
                'config': config_data
            })
        except Exception as e:
            return Response({
                'error': f'Config error: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    @action(detail=False, methods=['post'], url_path='upload-and-process')
    def upload_and_process(self, request):
        """
        UNIFIED ENDPOINT: Upload CRS file (PDF or Excel) and process it
        
        POST /api/v1/crs/documents/upload-and-process/
        Body (multipart/form-data):
            - file: PDF or Excel file (required)
            - project_name: string (optional)
            - document_number: string (optional)
            - revision: string (optional)
            - contractor: string (optional)
        
        Returns: Populated CRS template Excel file for download
        """
        if not HELPERS_AVAILABLE:
            return Response({
                'error': 'PDF/Excel processing helpers not available. Install PyPDF2 and openpyxl.',
                'success': False
            }, status=status.HTTP_503_SERVICE_UNAVAILABLE)
        
        try:
            # Get uploaded file
            uploaded_file = request.FILES.get('file')
            if not uploaded_file:
                return Response({
                    'error': 'File is required',
                    'success': False
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Check if preview mode is requested (returns JSON instead of file download)
            preview_mode = request.POST.get('preview', 'false').lower() == 'true'
            download_format = request.POST.get('format', 'xlsx').lower()  # xlsx, csv, json, pdf
            
            # Get metadata
            metadata = {
                'project_name': request.POST.get('project_name', ''),
                'document_number': request.POST.get('document_number', ''),
                'revision': request.POST.get('revision', ''),
                'contractor': request.POST.get('contractor', ''),
                'department': request.POST.get('department', ''),
            }
            
            # Determine file type
            file_name = uploaded_file.name.lower()
            is_pdf = file_name.endswith('.pdf')
            is_excel = file_name.endswith(('.xlsx', '.xls'))
            
            if not (is_pdf or is_excel):
                return Response({
                    'error': 'Invalid file type. Only PDF and Excel files are supported.',
                    'success': False
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Process based on file type
            if is_pdf:
                # Extract comments from PDF
                pdf_buffer = BytesIO(uploaded_file.read())
                comments = extract_reviewer_comments(pdf_buffer)
                
                if not comments:
                    return Response({
                        'error': 'No comments found in the PDF file.',
                        'success': False,
                        'message': 'The PDF does not contain any extractable reviewer comments.'
                    }, status=status.HTTP_400_BAD_REQUEST)
                
                # Get statistics
                stats = get_comment_statistics(comments)
                
                # Prepare comments data for preview
                comments_data = []
                for idx, comment in enumerate(comments, start=1):
                    comments_data.append({
                        'index': idx,
                        'reviewer_name': comment.reviewer_name,
                        'comment_text': comment.comment_text,
                        'discipline': comment.discipline,
                        'page_number': comment.page_number,
                        'comment_type': comment.comment_type,
                        'section_reference': comment.section_reference,
                        'raw_text': getattr(comment, 'raw_text', ''),
                        'cleaned': getattr(comment, 'cleaned', False),
                        'cleaning_method': getattr(comment, 'cleaning_method', ''),
                    })
                
                # PREVIEW MODE: Return JSON with extracted comments for user review
                if preview_mode:
                    return Response({
                        'success': True,
                        'message': f'Successfully extracted {len(comments)} comments',
                        'preview': True,
                        'metadata': metadata,
                        'statistics': {
                            'total': stats['total'],
                            'by_type': stats.get('by_type', {}),
                            'by_discipline': stats.get('by_discipline', {}),
                            'by_reviewer': stats.get('by_reviewer', {}),
                            'pages_with_comments': stats.get('pages_with_comments', 0),
                        },
                        'comments': comments_data,
                        'available_formats': ['xlsx', 'csv', 'json', 'pdf', 'docx'],
                    }, status=status.HTTP_200_OK)
                
                # DOWNLOAD MODE: Generate file in requested format
                # Load CRS template using smart template manager
                try:
                    template_buffer = get_crs_template()
                    # CRITICAL: Ensure buffer is at the start
                    template_buffer.seek(0)
                except FileNotFoundError as e:
                    return Response({
                        'error': 'CRS template not available',
                        'success': False,
                        'message': str(e),
                        'suggestion': 'Ensure CRS template is available locally or on AWS S3'
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Handle different download formats
                doc_number = metadata.get("document_number", "output")
                safe_filename = "".join(c for c in doc_number if c.isalnum() or c in "-_").strip() or "output"
                
                if download_format == 'json':
                    # Return as JSON file download
                    import json as json_module
                    json_content = json_module.dumps({
                        'metadata': metadata,
                        'statistics': stats,
                        'comments': comments_data
                    }, indent=2, default=str)
                    
                    response = HttpResponse(
                        json_content,
                        content_type='application/json'
                    )
                    response['Content-Disposition'] = f'attachment; filename="CRS_{safe_filename}.json"'
                    return response
                
                elif download_format == 'csv':
                    # Return as CSV file download
                    import csv
                    from io import StringIO
                    
                    csv_buffer = StringIO()
                    writer = csv.writer(csv_buffer)
                    
                    # Write header
                    writer.writerow(['No.', 'Reviewer', 'Comment', 'Discipline', 'Type', 'Page', 'Section'])
                    
                    # Write data
                    for c in comments_data:
                        writer.writerow([
                            c['index'],
                            c['reviewer_name'],
                            c['comment_text'],
                            c['discipline'],
                            c['comment_type'],
                            c['page_number'] or '',
                            c['section_reference']
                        ])
                    
                    response = HttpResponse(
                        csv_buffer.getvalue(),
                        content_type='text/csv'
                    )
                    response['Content-Disposition'] = f'attachment; filename="CRS_{safe_filename}.csv"'
                    return response
                
                elif download_format == 'pdf':
                    # Generate PDF report using ReportLab
                    from reportlab.lib import colors
                    from reportlab.lib.pagesizes import A4, landscape
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                    from reportlab.lib.enums import TA_LEFT, TA_CENTER
                    
                    pdf_buffer = BytesIO()
                    doc = SimpleDocTemplate(
                        pdf_buffer,
                        pagesize=landscape(A4),
                        rightMargin=30,
                        leftMargin=30,
                        topMargin=30,
                        bottomMargin=30
                    )
                    
                    elements = []
                    styles = getSampleStyleSheet()
                    
                    # Title
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Heading1'],
                        fontSize=16,
                        alignment=TA_CENTER,
                        spaceAfter=20
                    )
                    elements.append(Paragraph("Comment Resolution Sheet", title_style))
                    
                    # Metadata
                    meta_style = ParagraphStyle(
                        'Meta',
                        parent=styles['Normal'],
                        fontSize=10,
                        spaceAfter=5
                    )
                    if metadata.get('project_name'):
                        elements.append(Paragraph(f"<b>Project:</b> {metadata['project_name']}", meta_style))
                    if metadata.get('document_number'):
                        elements.append(Paragraph(f"<b>Document:</b> {metadata['document_number']}", meta_style))
                    if metadata.get('revision'):
                        elements.append(Paragraph(f"<b>Revision:</b> {metadata['revision']}", meta_style))
                    elements.append(Paragraph(f"<b>Total Comments:</b> {len(comments_data)}", meta_style))
                    elements.append(Spacer(1, 20))
                    
                    # Table data
                    table_data = [['#', 'Page', 'Type', 'Comment', 'Reviewer', 'Discipline']]
                    
                    cell_style = ParagraphStyle(
                        'Cell',
                        parent=styles['Normal'],
                        fontSize=8,
                        leading=10
                    )
                    
                    for c in comments_data:
                        # Truncate long comments for PDF
                        comment_text = c['comment_text'] or ''
                        if len(comment_text) > 200:
                            comment_text = comment_text[:200] + '...'
                        
                        table_data.append([
                            str(c['index']),
                            str(c['page_number'] or '-'),
                            c['comment_type'] or 'GENERAL',
                            Paragraph(comment_text, cell_style),
                            c['reviewer_name'] or '-',
                            c['discipline'] or '-'
                        ])
                    
                    # Create table
                    col_widths = [0.4*inch, 0.5*inch, 0.8*inch, 4.5*inch, 1.2*inch, 1*inch]
                    table = Table(table_data, colWidths=col_widths, repeatRows=1)
                    
                    table.setStyle(TableStyle([
                        # Header
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F46E5')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('TOPPADDING', (0, 0), (-1, 0), 8),
                        
                        # Body
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                        ('ALIGN', (1, 1), (1, -1), 'CENTER'),
                        ('ALIGN', (2, 1), (2, -1), 'CENTER'),
                        ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                        ('TOPPADDING', (0, 1), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                        
                        # Grid
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F3F4F6')]),
                    ]))
                    
                    elements.append(table)
                    doc.build(elements)
                    
                    pdf_buffer.seek(0)
                    pdf_content = pdf_buffer.read()
                    
                    response = HttpResponse(
                        pdf_content,
                        content_type='application/pdf'
                    )
                    response['Content-Disposition'] = f'attachment; filename="CRS_{safe_filename}.pdf"'
                    response['Content-Length'] = len(pdf_content)
                    return response
                
                elif download_format == 'docx':
                    # Generate Word document using python-docx
                    from docx import Document
                    from docx.shared import Inches, Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    from docx.oxml.ns import nsdecls
                    from docx.oxml import parse_xml
                    
                    doc = Document()
                    
                    # Title
                    title = doc.add_heading('Comment Resolution Sheet', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Metadata section
                    if metadata.get('project_name'):
                        doc.add_paragraph(f"Project: {metadata['project_name']}")
                    if metadata.get('document_number'):
                        doc.add_paragraph(f"Document Number: {metadata['document_number']}")
                    if metadata.get('revision'):
                        doc.add_paragraph(f"Revision: {metadata['revision']}")
                    if metadata.get('contractor'):
                        doc.add_paragraph(f"Contractor: {metadata['contractor']}")
                    if metadata.get('department'):
                        doc.add_paragraph(f"Department: {metadata['department']}")
                    doc.add_paragraph(f"Total Comments: {len(comments_data)}")
                    doc.add_paragraph()  # Empty line
                    
                    # Create table
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    headers = ['#', 'Page', 'Type', 'Comment', 'Reviewer', 'Discipline']
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        # Make header bold
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                        # Header background color (indigo)
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4F46E5"/>')
                        header_cells[i]._tc.get_or_add_tcPr().append(shading)
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    # Data rows
                    for c in comments_data:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(c['index'])
                        row_cells[1].text = str(c['page_number'] or '-')
                        row_cells[2].text = c['comment_type'] or 'GENERAL'
                        row_cells[3].text = c['comment_text'] or '-'
                        row_cells[4].text = c['reviewer_name'] or '-'
                        row_cells[5].text = c['discipline'] or '-'
                    
                    # Set column widths
                    widths = [Inches(0.4), Inches(0.5), Inches(0.8), Inches(4.0), Inches(1.2), Inches(1.0)]
                    for row in table.rows:
                        for idx, cell in enumerate(row.cells):
                            cell.width = widths[idx]
                    
                    # Save to buffer
                    docx_buffer = BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    docx_content = docx_buffer.read()
                    
                    response = HttpResponse(
                        docx_content,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    response['Content-Disposition'] = f'attachment; filename="CRS_{safe_filename}.docx"'
                    response['Content-Length'] = len(docx_content)
                    return response
                
                else:  # xlsx (default)
                    # Populate template
                    output_buffer = populate_crs_template(
                        template_buffer,
                        comments,
                        metadata
                    )
                
                # CRITICAL: Ensure buffer is at the beginning before reading
                output_buffer.seek(0)
                output_content = output_buffer.read()
                
                # Validate output is not empty
                if len(output_content) < 100:
                    return Response({
                        'error': 'Failed to generate valid Excel file',
                        'success': False,
                        'message': 'Output file is too small, likely corrupted'
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Create response with Excel file
                response = HttpResponse(
                    output_content,
                    content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                )
                
                response['Content-Disposition'] = f'attachment; filename="CRS_Populated_{safe_filename}.xlsx"'
                response['Content-Length'] = len(output_content)
                response['X-Comment-Count'] = str(stats['total'])
                response['X-Processing-Status'] = 'success'
                
                return response
            
            elif is_excel:
                # Excel file uploaded - validate and standardize format
                return Response({
                    'error': 'Excel processing coming soon',
                    'success': False,
                    'message': 'Please upload a PDF file for now. Excel input will be supported soon.'
                }, status=status.HTTP_501_NOT_IMPLEMENTED)
            
        except Exception as e:
            import traceback
            return Response({
                'error': f'Error processing file: {str(e)}',
                'success': False,
                'details': traceback.format_exc()
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class CRSCommentViewSet(viewsets.ModelViewSet):
    """ViewSet for CRS Comment management"""
    queryset = CRSComment.objects.all()
    serializer_class = CRSCommentSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        """Filter by document_id if provided"""
        queryset = super().get_queryset()
        document_id = self.request.query_params.get('document_id')
        if document_id:
            queryset = queryset.filter(document_id=document_id)
        return queryset
    
    def perform_update(self, serializer):
        """Log update and recalculate document stats"""
        old_instance = self.get_object()
        old_status = old_instance.status
        
        serializer.save()
        
        new_instance = serializer.instance
        document = new_instance.document
        
        # Log status change
        if old_status != new_instance.status:
            CRSActivity.objects.create(
                document=document,
                comment=new_instance,
                action='status_changed',
                description=f'Comment #{new_instance.serial_number} status changed from {old_status} to {new_instance.status}',
                performed_by=self.request.user,
                old_value={'status': old_status},
                new_value={'status': new_instance.status}
            )
            
            # Set resolved_at timestamp
            if new_instance.status in ['resolved', 'closed'] and not new_instance.resolved_at:
                new_instance.resolved_at = timezone.now()
                new_instance.save()
        
        # Recalculate document stats
        self.update_document_stats(document)
    
    def update_document_stats(self, document):
        """Update document statistics based on comments"""
        total = document.comments.count()
        resolved = document.comments.filter(status__in=['resolved', 'closed']).count()
        pending = total - resolved
        
        document.total_comments = total
        document.resolved_comments = resolved
        document.pending_comments = pending
        
        # Update status
        if total > 0 and resolved == total:
            document.status = 'completed'
        elif resolved > 0:
            document.status = 'in_review'
        
        document.save()
    
    @action(detail=True, methods=['post'])
    def add_contractor_response(self, request, pk=None):
        """
        Add contractor response to comment
        POST /api/v1/crs/comments/{id}/add_contractor_response/
        Body: { "response": "text" }
        """
        comment = self.get_object()
        response_text = request.data.get('response', '').strip()
        
        if not response_text:
            return Response(
                {"error": "Response text is required"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        comment.contractor_response = response_text
        comment.contractor_response_date = timezone.now()
        comment.contractor_responder = request.user
        
        if comment.status == 'open':
            comment.status = 'in_progress'
        
        comment.save()
        
        # Log activity
        CRSActivity.objects.create(
            document=comment.document,
            comment=comment,
            action='response_added',
            description=f'Contractor response added to comment #{comment.serial_number}',
            performed_by=request.user
        )
        
        return Response({
            "success": True,
            "message": "Contractor response added successfully",
            "data": CRSCommentSerializer(comment).data
        })
    
    @action(detail=True, methods=['post'])
    def add_company_response(self, request, pk=None):
        """
        Add company response to comment
        POST /api/v1/crs/comments/{id}/add_company_response/
        Body: { "response": "text" }
        """
        comment = self.get_object()
        response_text = request.data.get('response', '').strip()
        
        if not response_text:
            return Response(
                {"error": "Response text is required"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        comment.company_response = response_text
        comment.company_response_date = timezone.now()
        comment.company_responder = request.user
        
        # If both responses exist, mark as resolved
        if comment.contractor_response and comment.company_response:
            comment.status = 'resolved'
            comment.resolved_at = timezone.now()
        elif comment.status == 'open':
            comment.status = 'in_progress'
        
        comment.save()
        
        # Update document stats
        self.update_document_stats(comment.document)
        
        # Log activity
        CRSActivity.objects.create(
            document=comment.document,
            comment=comment,
            action='response_added',
            description=f'Company response added to comment #{comment.serial_number}',
            performed_by=request.user
        )
        
        return Response({
            "success": True,
            "message": "Company response added successfully",
            "data": CRSCommentSerializer(comment).data
        })
    
    @action(detail=False, methods=['post'])
    def bulk_create(self, request):
        """
        Bulk create comments
        POST /api/v1/crs/comments/bulk_create/
        Body: { "document_id": 1, "comments": [{...}, {...}] }
        """
        serializer = CRSCommentBulkCreateSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        document_id = serializer.validated_data['document_id']
        comments_data = serializer.validated_data['comments']
        
        try:
            document = CRSDocument.objects.get(id=document_id)
        except CRSDocument.DoesNotExist:
            return Response(
                {"error": "Document not found"},
                status=status.HTTP_404_NOT_FOUND
            )
        
        with transaction.atomic():
            created_comments = []
            for idx, comment_data in enumerate(comments_data, start=1):
                comment = CRSComment.objects.create(
                    document=document,
                    serial_number=idx,
                    page_number=comment_data.get('page', 1),
                    clause_number=comment_data.get('clause', ''),
                    comment_text=comment_data.get('text', ''),
                    comment_type=comment_data.get('type', 'other'),
                    status='open'
                )
                created_comments.append(comment)
            
            # Update document stats
            document.total_comments = len(created_comments)
            document.pending_comments = len(created_comments)
            document.resolved_comments = 0
            document.save()
        
        return Response({
            "success": True,
            "message": f"Created {len(created_comments)} comments",
            "data": CRSCommentSerializer(created_comments, many=True).data
        }, status=status.HTTP_201_CREATED)


class CRSActivityViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for CRS Activity logs (read-only)"""
    queryset = CRSActivity.objects.all()
    serializer_class = CRSActivitySerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        """Filter by document_id if provided"""
        queryset = super().get_queryset()
        document_id = self.request.query_params.get('document_id')
        if document_id:
            queryset = queryset.filter(document_id=document_id)
        return queryset


class GoogleSheetConfigViewSet(viewsets.ModelViewSet):
    """ViewSet for Google Sheet Configuration"""
    queryset = GoogleSheetConfig.objects.all()
    serializer_class = GoogleSheetConfigSerializer
    permission_classes = [IsAuthenticated]
    
    def perform_create(self, serializer):
        """Set created_by to current user"""
        serializer.save(created_by=self.request.user)
    
    @action(detail=True, methods=['post'])
    def test_connection(self, request, pk=None):
        """
        Test Google Sheets connection
        POST /api/v1/crs/google-configs/{id}/test_connection/
        """
        config = self.get_object()
        
        try:
            gs_service = GoogleSheetsService(
                credentials_json=config.credentials_json,
                token_json=config.token_json
            )
            
            if gs_service.authenticate():
                # Save updated token
                new_token = gs_service.get_token_json()
                if new_token:
                    config.token_json = new_token
                    config.save()
                
                return Response({
                    "success": True,
                    "message": "Successfully connected to Google Sheets API"
                })
            else:
                return Response({
                    "success": False,
                    "message": "Failed to authenticate with Google Sheets API"
                }, status=status.HTTP_400_BAD_REQUEST)
        
        except Exception as e:
            return Response({
                "success": False,
                "message": f"Error testing connection: {str(e)}"
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
